"""Document I/O for reading and writing DOCX files."""

import logging
from typing import Dict, Any
from pathlib import Path

from docx import Document
from utils import DocumentAnalyzer, PathResolver

logger = logging.getLogger(__name__)


class DocumentReader:
    """Reads DOCX files and extracts structured content including images."""

    def _run(self, file_path: str) -> Dict[str, Any]:
        """Read and analyze document structure including images."""
        try:
            path = PathResolver.resolve_file_path(file_path)
            logger.info(f"Reading document from: {path}")

            doc = Document(str(path))

            content = {
                'paragraphs': [],
                'tables': [],
                'images': {},
                'metadata': {
                    'total_paragraphs': 0,
                    'total_tables': len(doc.tables),
                    'total_images': 0,
                    'has_footnotes': False
                }
            }

            content['images'] = DocumentAnalyzer.extract_images_from_document(doc)
            content['metadata']['total_images'] = len(content['images'])

            for i, paragraph in enumerate(doc.paragraphs):
                structure = DocumentAnalyzer.analyze_paragraph_structure(paragraph)
                structure['index'] = i
                if paragraph.text.strip() or structure['has_image']:
                    content['paragraphs'].append(structure)

            content['tables'] = DocumentAnalyzer.extract_tables_from_document(doc)
            content['metadata']['total_paragraphs'] = len(content['paragraphs'])

            logger.info(
                f"Successfully analyzed document: {len(content['paragraphs'])} paragraphs, "
                f"{len(content['tables'])} tables, {len(content['images'])} images"
            )
            return content

        except Exception as e:
            logger.error(f"Error reading document: {str(e)}")
            raise


class DocumentWriter:
    """Writes translated content to DOCX with original formatting and images."""

    def _run(self, translated_content: Dict[str, Any], output_path: str,
             original_doc_path: str = None) -> str:
        """Write translated content to DOCX file preserving images."""
        try:
            if original_doc_path and Path(original_doc_path).exists():
                doc = self._update_original_document(translated_content, original_doc_path)
            else:
                doc = self._create_new_document(translated_content)

            doc.save(output_path)
            logger.info(f"Document saved to: {output_path}")
            return f"Document successfully saved to {output_path}"

        except Exception as e:
            logger.error(f"Error writing document: {str(e)}")
            raise

    def _update_original_document(self, translated_content: Dict[str, Any],
                                  original_doc_path: str) -> Document:
        """Update original document with translations while preserving formatting and images."""
        doc = Document(original_doc_path)
        logger.info("Using original document as template to preserve images and formatting")

        translated_paragraphs = {
            p.get('index', i): p
            for i, p in enumerate(translated_content.get('paragraphs', []))
        }

        for i, paragraph in enumerate(doc.paragraphs):
            if i in translated_paragraphs:
                self._update_paragraph_text(paragraph, translated_paragraphs[i])

        return doc

    def _update_paragraph_text(self, paragraph, trans_para: Dict[str, Any]):
        """Update paragraph text while preserving images and run-level formatting."""
        translated_text = trans_para.get('text', '').strip()
        if not translated_text:
            return

        if trans_para.get('has_image', False):
            # For mixed image+text paragraphs update the first text-bearing run
            for run in paragraph.runs:
                if not run._element.xpath('.//pic:pic') and run.text.strip():
                    run.text = translated_text
                    return
            paragraph.add_run(" " + translated_text)
            return

        # Collect all non-image runs
        text_runs = [
            run for run in paragraph.runs
            if not run._element.xpath('.//pic:pic')
        ]

        if not text_runs:
            paragraph.add_run(translated_text)
            return

        # Use the run with the most text as the carrier so its formatting
        # (bold heading, italic caption, etc.) is the dominant style applied
        # to the full translated text.
        dominant_run = max(text_runs, key=lambda r: len(r.text))

        # Clear all other text runs
        for run in text_runs:
            if run is not dominant_run:
                run.text = ""

        dominant_run.text = translated_text

    def _create_new_document(self, translated_content: Dict[str, Any]) -> Document:
        """Create new document if original not available (images will not be preserved)."""
        doc = Document()
        logger.info("Creating new document (images may not be preserved)")

        for para_data in translated_content.get('paragraphs', []):
            p = doc.add_paragraph()

            if para_data.get('text', '').strip():
                run = p.add_run(para_data['text'])
                if para_data.get('is_bold', False):
                    run.bold = True
                if para_data.get('is_italic', False):
                    run.italic = True

            if para_data.get('alignment'):
                p.alignment = para_data['alignment']

            if para_data.get('has_image', False):
                logger.warning(
                    f"Image in paragraph {para_data.get('index', 'unknown')} "
                    "cannot be preserved without the original document"
                )

        for table_data in translated_content.get('tables', []):
            if table_data.get('data'):
                table = doc.add_table(
                    rows=table_data['rows'],
                    cols=table_data['cols']
                )
                for row_idx, row_data in enumerate(table_data['data']):
                    for col_idx, cell_text in enumerate(row_data):
                        if row_idx < len(table.rows) and col_idx < len(table.rows[0].cells):
                            table.rows[row_idx].cells[col_idx].text = cell_text

        return doc
